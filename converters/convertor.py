import fitz
from docx import Document


class FileConverter:

    def pdf_to_txt(self, input_file, output_file):

        pdf = fitz.open(input_file)

        text = ""

        for page in pdf:
            text += page.get_text()

        pdf.close()

        with open(output_file, "w", encoding="utf-8") as file:
            file.write(text)

        print("PDF → TXT completed!")

    def docx_to_txt(self, input_file, output_file):

        doc = Document(input_file)

        text = ""

        for para in doc.paragraphs:
            text += para.text + "\n"

        with open(output_file, "w", encoding="utf-8") as file:
            file.write(text)

        print("DOCX → TXT completed!")

    def txt_to_docx(self, input_file, output_file):

        with open(input_file, "r", encoding="utf-8") as file:
            text = file.read()

        doc = Document()

        doc.add_paragraph(text)

        doc.save(output_file)

        print("TXT → DOCX completed!")

    def convert(self, input_file, output_file):

        input_ext = input_file.split(".")[-1].lower()

        output_ext = output_file.split(".")[-1].lower()

        conversion_map = {

            ("pdf", "txt"): self.pdf_to_txt,

            ("docx", "txt"): self.docx_to_txt,

            ("txt", "docx"): self.txt_to_docx
        }

        key = (input_ext, output_ext)

        if key in conversion_map:

            conversion_map[key](input_file, output_file)

        else:

            print("Conversion not supported!")


# MAIN

converter = FileConverter()

input_file = input("Enter input file: ")

output_file = input("Enter output file: ")

converter.convert(input_file, output_file)